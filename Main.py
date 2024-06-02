import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import openpyxl
from docx import Document
import os

input_folder_path = None

def folder_selection():
    global input_folder_path
    
    # Önceki dosyaya ait verileri silme işlemi ve onay mesajı
    if input_folder_path:
        confirm = messagebox.askyesno("Uyarı", "Önceki dosyada yapılan değişiklikler silinecek. Devam etmek istiyor musunuz?")
        if not confirm:
            return
        else:
            treeview.delete(*treeview.get_children())

    # Yeni dosya seçimi
    new_input_folder_path = filedialog.askopenfilename(filetypes=[("Excel Documents", "*.xlsx;*.xlsm")])
    
    if new_input_folder_path:
        input_folder_path = new_input_folder_path
        print("Selected file:", input_folder_path)
        load_data(input_folder_path)


def on_focus_in(event, placeholder):
    entry = event.widget
    if entry.get() == placeholder:
        entry.delete(0, tk.END)

def on_focus_out(event, placeholder):
    entry = event.widget
    if entry.get() == "":
        entry.insert(0, placeholder)

def list_filter():
    filter_text = dispo_combobox.get()
    if not filter_text:
        # Filtre seçilmemişse, tüm veriyi göster
        for row_id in treeview.get_children():
            treeview.item(row_id, tags=())
    else:
        # Seçilen filtreye göre veriyi filtrele
        for row_id in treeview.get_children():
            values = treeview.item(row_id, 'values')
            disposition_type = values[1]  # Disposition Type sütununun indeksi
            if disposition_type == filter_text:
                treeview.item(row_id, tags=('visible',))
            else:
                treeview.item(row_id, tags=('hidden',))
    
    # Filtrelenen veriyi göster
    treeview.tag_configure('visible', background='green')
    treeview.tag_configure('hidden', background='grey')

def undo_filter():
        for row_id in treeview.get_children():
            treeview.item(row_id, tags=())


def print_karar():
    # Seçilen disposition tipi
    disposition_type = dispo_combobox.get()
    
    # Excel dosyasındaki verileri oku
    dimensional_flaws = []
    serial_number = ""
    operation_number = ""
    cause_op_number = ""
    quantity = ""
    cause_code = ""
    
    if input_folder_path:
        workbook = openpyxl.load_workbook(input_folder_path)
        sheet = workbook["DISPOSITION"]
        for row in sheet.iter_rows(min_row=2, max_col=7, max_row=sheet.max_row):
            if row[1].value == disposition_type:  # "Disposition Type" sütunu ile eşleşen satırları bul
                dimensional_flaws.append(str(row[0].value))  # "Dimensional Flaw" değerlerini listeye ekle
                serial_number = row[2].value
                operation_number = row[3].value
                cause_op_number = row[4].value
                quantity = row[5].value
                cause_code = row[6].value
    
    # Dosya adını oluşturma
    new_filename = f"{serial_number}_{disposition_type}.docx"
    
    # Masaüstünde dosya arama ve placeholderları değiştirme
    desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
    for filename in os.listdir(desktop_path):
        if filename.endswith(".docx") and disposition_type in filename:
            docx_path = os.path.join(desktop_path, filename)
            document = Document(docx_path)
            
            # Placeholder'ları değiştirme
            placeholders = {
                "[Serial Number]": serial_number,
                "[Operation Number]": operation_number,
                "[Cause OP. Number]": cause_op_number,
                "[Dimensional Flaw]": dimensional_flaws,  # "Dimensional Flaw" değerlerinin listesi
                "[Quantitiy]": quantity,
                "[Cause Code]": cause_code,
            }
            
            # Paragrafları kontrol etme
            for paragraph in document.paragraphs:
                for placeholder, value in placeholders.items():
                    if placeholder in paragraph.text:
                        if isinstance(value, list):  # "Dimensional Flaw" için
                            if value:  # Liste boş değilse
                                paragraph.text = paragraph.text.replace(placeholder, value.pop(0))  # Listenin ilk elemanını al ve placeholder ile değiştir
                        else:
                            paragraph.text = paragraph.text.replace(placeholder, str(value))
            
            # Tabloları kontrol etme
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for placeholder, value in placeholders.items():
                                if placeholder in paragraph.text:
                                    if isinstance(value, list):  # "Dimensional Flaw" için
                                        if value:  # Liste boş değilse
                                            paragraph.text = paragraph.text.replace(placeholder, value.pop(0))  # Listenin ilk elemanını al ve placeholder ile değiştir
                                    else:
                                        paragraph.text = paragraph.text.replace(placeholder, str(value))
            
            document.save(os.path.join(desktop_path, new_filename))
            
            # Başarı mesajı göster
            messagebox.showinfo("Bilgi", "Belge başarıyla oluşturuldu.")
            break
    else:
        messagebox.showerror("Hata", f"{disposition_type} ile başlayan bir docx dosyası bulunamadı.")
    





def focus_deneme():
    selected_items = treeview.selection()
    if not selected_items:
        messagebox.showinfo("Bilgilendirme", "Lütfen en az bir satır seçin.")
        return
    
    name = name_entry.get()
    operation = operation_entry.get()
    cause_operation = cause_operation_entry.get()
    quantity = qty_entry.get()
    cause_code = ccode_entry.get()
    disposition = status_combobox.get()
    
    path = input_folder_path
    workbook = openpyxl.load_workbook(path)
    sheet = workbook["DISPOSITION"]

    for item in selected_items:
        item_values = treeview.item(item, 'values')
        row_id = int(item[1:])  # Treeview item IDs start with 'I' followed by a number

        sheet.cell(row=row_id+1, column=2).value = disposition
        sheet.cell(row=row_id+1, column=3).value = name
        sheet.cell(row=row_id+1, column=4).value = operation
        sheet.cell(row=row_id+1, column=5).value = cause_operation
        sheet.cell(row=row_id+1, column=6).value = quantity
        sheet.cell(row=row_id+1, column=7).value = cause_code
        
        # Update Treeview
        new_values = list(item_values)
        new_values[1] = disposition
        new_values[2] = name
        new_values[3] = operation
        new_values[4] = cause_operation
        new_values[5] = quantity
        new_values[6] = cause_code
        treeview.item(item, values=new_values)
    
    workbook.save(path)
    
    status_combobox.set("")
    name_entry.delete(0, "end")
    name_entry.insert(0, "Serial Number")
    operation_entry.delete(0, "end")
    operation_entry.insert(0, "Operation Number")
    cause_operation_entry.delete(0, "end")
    cause_operation_entry.insert(0, "Cause OP. Number")
    qty_entry.delete(0, "end")
    qty_entry.insert(0, "Quantity")
    ccode_entry.delete(0, "end")
    ccode_entry.insert(0, "Cause Code")
    
    for item in selected_items:
        treeview.selection_remove(item)
    
    print("Seçim Kaldırıldı.")
    messagebox.showinfo("Bilgilendirme", "Seçilen satırlar için atama yapıldı")

def insert_row():
    name = name_entry.get()
    subscription_status = status_combobox.get()
    path = input_folder_path
    workbook = openpyxl.load_workbook(path)
    sheet = workbook.active
    row_values = [name, subscription_status]
    sheet.append(row_values)
    workbook.save(path)
    treeview.insert("", tk.END, values=row_values)
    name_entry.delete(0, "end")
    name_entry.insert(0, "Serial Number")
    status_combobox.set(combo_list[0])

def load_data(filepath):
    workbook = openpyxl.load_workbook(filepath)
    if "DISPOSITION" in workbook.sheetnames:
        sheet = workbook["DISPOSITION"]
        list_values = list(sheet.values)
        cols = list_values[0]
        treeview.config(columns=cols)
        for col in cols:
            treeview.heading(col, text=col)
            treeview.column(col, anchor='center')
        for value_tuple in list_values[1:]:
            treeview.insert("", tk.END, values=value_tuple)
    else:
        messagebox.showerror("Error", "The 'DISPOSITION' sheet does not exist in the selected file.")

root = tk.Tk()
root.title("Atos Disposition Creator")

style = ttk.Style(root)
root.tk.call("source", "forest-dark.tcl")
root.tk.call("source", "forest-light.tcl")
style.theme_use("forest-dark")

frame = ttk.Frame(root)
frame.pack(fill=tk.BOTH, expand=True)

root.columnconfigure(0, weight=1)
root.rowconfigure(0, weight=1)

frame.columnconfigure(0, weight=1)
frame.columnconfigure(1, weight=1)
frame.columnconfigure(2, weight=1)
frame.rowconfigure(0, weight=1)
frame.rowconfigure(1, weight=1)
frame.rowconfigure(2, weight=1)
frame.rowconfigure(3, weight=1)

# Dosya Seçme Frame
dosya_frame = ttk.LabelFrame(frame, text='Rapor Dosyası Seçimi')
dosya_frame.grid(row=0, column=0, padx=20, pady=10, sticky='nsew')

folder_button = ttk.Button(dosya_frame, text="Dosya Seçimi", command=folder_selection)
folder_button.grid(row=0, column=0, padx=5, pady=5, sticky='nsew')

# Karar Frame
karar_frame = ttk.LabelFrame(frame, text='Karar Atanması')
karar_frame.grid(row=1, column=0, padx=20, pady=10, sticky='nsew')

karar_frame.columnconfigure(0, weight=1)
karar_frame.columnconfigure(1, weight=1)
karar_frame.rowconfigure(0, weight=1)
karar_frame.rowconfigure(1, weight=1)
karar_frame.rowconfigure(2, weight=1)
karar_frame.rowconfigure(3, weight=1)
karar_frame.rowconfigure(4, weight=1)

serial_number_placeholder = "Serial Number"
name_entry = ttk.Entry(karar_frame)
name_entry.insert(0, serial_number_placeholder)
name_entry.bind("<FocusIn>", lambda e: on_focus_in(e, serial_number_placeholder))
name_entry.bind("<FocusOut>", lambda e: on_focus_out(e, serial_number_placeholder))
name_entry.grid(row=0, column=0, padx=5, pady=(0,5), sticky="ew")

operation_entry_placeholder = "Operation Number"
operation_entry = ttk.Entry(karar_frame)
operation_entry.insert(0, operation_entry_placeholder)
operation_entry.bind("<FocusIn>", lambda e: on_focus_in(e, operation_entry_placeholder))
operation_entry.bind("<FocusOut>", lambda e: on_focus_out(e, operation_entry_placeholder))
operation_entry.grid(row=0, column=1, padx=5, pady=(0,5), sticky="ew")

cause_operation_entry_placeholder = "Cause OP. Number"
cause_operation_entry = ttk.Entry(karar_frame)
cause_operation_entry.insert(0, cause_operation_entry_placeholder)
cause_operation_entry.bind("<FocusIn>", lambda e: on_focus_in(e, cause_operation_entry_placeholder))
cause_operation_entry.bind("<FocusOut>", lambda e: on_focus_out(e, cause_operation_entry_placeholder))
cause_operation_entry.grid(row=1, column=0, padx=5, pady=(0,5), sticky="ew")

qty_entry_placeholder = "Quantity"
qty_entry = ttk.Entry(karar_frame)
qty_entry.insert(0, qty_entry_placeholder)
qty_entry.bind("<FocusIn>", lambda e: on_focus_in(e, qty_entry_placeholder))
qty_entry.bind("<FocusOut>", lambda e: on_focus_out(e, qty_entry_placeholder))
qty_entry.grid(row=1, column=1, padx=5, pady=(0,5), sticky="ew")

ccode_entry_placeholder = "Cause Code"
ccode_entry = ttk.Entry(karar_frame)
ccode_entry.insert(0, ccode_entry_placeholder)
ccode_entry.bind("<FocusIn>", lambda e: on_focus_in(e, ccode_entry_placeholder))
ccode_entry.bind("<FocusOut>", lambda e: on_focus_out(e, ccode_entry_placeholder))
ccode_entry.grid(row=2, column=0, padx=5, pady=(0,5), sticky="ew")

combo_list = ["", "MRB", "DEBURR R/W", "CTP & R/I", "CTP & R/W", "STD OP. R/W", "ACCEPT", "RETURN TO VENDOR"]
combo_list.sort()

status_combobox = ttk.Combobox(karar_frame, values=combo_list, state="readonly")
status_combobox.grid(row=2, column=1, padx=5, pady=5, sticky='ew')

separator = ttk.Separator(karar_frame)
separator.grid(row=3, column=0, padx=(20,10), pady=10, sticky="ew")

button = ttk.Button(karar_frame, text="Insert", command=focus_deneme)
button.grid(row=4, column=0, padx=5, pady=5, sticky='nsew')

separator = ttk.Separator(karar_frame)
separator.grid(row=5, column=0, padx=(20,10), pady=10, sticky="ew")

# Yazdırma Kolonu
dispo_frame = ttk.LabelFrame(frame, text='Karar Yazdırma Kolonu')
dispo_frame.grid(row=1, column=1, columnspan=2, padx=20, pady=10, sticky='nsew')

dispo_frame.columnconfigure(0, weight=1)
dispo_frame.columnconfigure(1, weight=1)
dispo_frame.rowconfigure(0, weight=1)
dispo_frame.rowconfigure(1, weight=1)
dispo_frame.rowconfigure(2, weight=1)
dispo_frame.rowconfigure(3, weight=1)
dispo_frame.rowconfigure(4, weight=1)
dispo_frame.rowconfigure(5, weight=1)
dispo_frame.rowconfigure(6, weight=1)

dispo_list = ["", "MRB", "DEBURR R/W", "CTP & R/I", "CTP & R/W", "STD OP. R/W", "ACCEPT", "RETURN TO VENDOR", "WELD R/W"]
dispo_list.sort()

dispo_combobox = ttk.Combobox(dispo_frame, values=dispo_list, state="readonly")
dispo_combobox.grid(row=0, column=1, padx=5, pady=5, sticky='ew')

separator = ttk.Separator(dispo_frame)
separator.grid(row=1, column=1, padx=(20,10), pady=10, sticky="ew")

button_filter = ttk.Button(dispo_frame, text="Filtrele", command=list_filter)
button_filter.grid(row=2, column=1, padx=5, pady=5, sticky='nsew')

separator = ttk.Separator(dispo_frame)
separator.grid(row=3, column=1, padx=(20,10), pady=10, sticky="ew")

button_filter_undo = ttk.Button(dispo_frame, text="Filtre Kaldır", command=undo_filter)
button_filter_undo.grid(row=4, column=1, padx=5, pady=5, sticky='nsew')

separator = ttk.Separator(dispo_frame)
separator.grid(row=5, column=1, padx=(20,10), pady=10, sticky="ew")

button_print = ttk.Button(dispo_frame, text="Yazdır", command=print_karar)
button_print.grid(row=6, column=1, padx=5, pady=5, sticky='nsew')

# TreeView
treeFrame = ttk.Frame(frame)
treeFrame.grid(row=3, column=0, columnspan=3, pady=10, sticky='nsew')
treeFrame.columnconfigure(0, weight=1)
treeFrame.rowconfigure(0, weight=1)

treeScroll = ttk.Scrollbar(treeFrame)
treeScroll.pack(side="right", fill="y")

treeview = ttk.Treeview(treeFrame, show="headings", yscrollcommand=treeScroll.set, height=13)
treeview.pack(fill=tk.BOTH, expand=True)
treeScroll.config(command=treeview.yview)

# Yatay Scrollbar oluştur
treeScrollX = ttk.Scrollbar(treeFrame, orient="horizontal", command=treeview.xview)
treeScrollX.pack(side="bottom", fill="x")

# Treeview'e yatay Scrollbar bağla
treeview.configure(xscrollcommand=treeScrollX.set)

root.mainloop()
